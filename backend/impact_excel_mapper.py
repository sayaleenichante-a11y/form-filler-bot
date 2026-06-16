import csv
import re
import zipfile;
from difflib import SequenceMatcher
from pathlib import Path

import openpyxl

from format_utils import (
    clean_value,
    format_value_for_field,
    normalize_text
)


try:
    import fitz
except ImportError:
    fitz = None


try:
    from docx import Document
except ImportError:
    Document = None


def similarity(a, b):
    a = normalize_text(a)
    b = normalize_text(b)

    if not a or not b:
        return 0.0

    if a == b:
        return 1.0

    if a in b or b in a:
        return 0.92

    a_words = set(a.split())
    b_words = set(b.split())

    token_score = (
        len(a_words & b_words)
        / max(
            len(a_words),
            len(b_words),
            1
        )
    )

    sequence_score = SequenceMatcher(
        None,
        a,
        b
    ).ratio()

    return max(
        token_score,
        sequence_score
    )


TERM_GROUPS = [
    {
        "name",
        "full name",
        "applicant name",
        "candidate name",
        "employee name",
        "student name"
    },
    {
        "dob",
        "date of birth",
        "birth date"
    },
    {
        "phone",
        "mobile",
        "contact number",
        "telephone"
    },
    {
        "email",
        "email address",
        "mail id"
    },
    {
        "company",
        "organisation",
        "organization",
        "firm"
    },
    {
        "designation",
        "job title",
        "position",
        "role"
    },
    {
        "salary",
        "ctc",
        "package",
        "stipend",
        "income"
    },
    {
        "degree",
        "qualification",
        "course",
        "program"
    },
    {
        "branch",
        "stream",
        "specialization",
        "specialisation"
    },
    {
        "percentage",
        "percent",
        "marks"
    },
    {
        "cgpa",
        "gpa",
        "pointer"
    },
    {
        "address",
        "residential address",
        "permanent address",
        "current address"
    },
    {
        "city",
        "location",
        "town"
    },
    {
        "pincode",
        "pin code",
        "postal code",
        "zip code"
    },
    {
        "date",
        "application date",
        "submission date"
    },
    {
        "description",
        "details",
        "remarks",
        "summary"
    },
]


def expanded_terms(text):
    normalized = normalize_text(text)
    values = {normalized}

    for group in TERM_GROUPS:
        normalized_group = {
            normalize_text(item)
            for item in group
        }

        matched = normalized in normalized_group

        if not matched:
            for item in normalized_group:
                if (
                    len(
                        item.replace(" ", "")
                    ) >= 5
                    and (
                        item in normalized
                        or normalized in item
                    )
                ):
                    matched = True
                    break

        if matched:
            values.update(
                normalized_group
            )

    return {
        item
        for item in values
        if item
    }


def semantic_score(a, b):
    scores = [
        similarity(first, second)
        for first in expanded_terms(a)
        for second in expanded_terms(b)
    ]

    return max(
        scores,
        default=0.0
    )


def is_empty_row(row):
    return not any(
        clean_value(value)
        for value in row
    )


def _is_plain_number(value):
    text = clean_value(value).replace(",", "")

    return bool(
        re.fullmatch(
            r"-?\d+(?:\.\d+)?",
            text
        )
    )


def header_row_score(row):
    values = [
        clean_value(value)
        for value in row
        if clean_value(value)
    ]

    if len(values) < 2:
        return -1

    alphabetic_count = sum(
        1
        for value in values
        if re.search(r"[A-Za-z]", value)
    )

    numeric_count = sum(
        1
        for value in values
        if _is_plain_number(value)
    )

    unique_count = len({
        normalize_text(value)
        for value in values
    })

    return (
        alphabetic_count * 2
        + unique_count
        - numeric_count
    )


def detect_header_row(matrix):
    best_index = 0
    best_score = -1

    for index, row in enumerate(matrix[:30]):
        score = header_row_score(row)

        if score > best_score:
            best_score = score
            best_index = index

    return best_index


def _is_probable_header_continuation(row):
    values = [
        clean_value(value)
        for value in row
        if clean_value(value)
    ]

    if len(values) < 2:
        return False

    # Data row me numbers hote hain.
    # Second/third heading row me generally labels hote hain.
    if any(
        _is_plain_number(value)
        for value in values
    ):
        return False

    alphabetic_count = sum(
        1
        for value in values
        if re.search(r"[A-Za-z]", value)
    )

    return alphabetic_count >= max(
        2,
        int(len(values) * 0.60)
    )


def _row_signature(row, width):
    return tuple(
        normalize_text(
            row[index]
            if index < len(row)
            else ""
        )
        for index in range(width)
    )


def make_unique_headers(headers):
    result = []
    used = {}

    for index, header in enumerate(headers):
        base = (
            clean_value(header)
            or f"Column {index + 1}"
        )

        key = normalize_text(base)

        count = used.get(key, 0) + 1
        used[key] = count

        result.append(
            base
            if count == 1
            else f"{base} ({count})"
        )

    return result


def _build_multi_row_headers(
    header_rows,
    width
):
    rows = [
        [
            clean_value(
                row[column]
                if column < len(row)
                else ""
            )
            for column in range(width)
        ]
        for row in header_rows
    ]

    # Merged parent headings ko child columns par
    # automatically forward-fill karega.
    #
    # Example:
    # Lat + Deg/Min/Sec/Dir
    # Long + Deg/Min/Sec/Dir
    for row_index in range(
        len(rows) - 1
    ):
        last_parent = ""

        for column in range(width):
            current_value = (
                rows[row_index][column]
            )

            if current_value:
                last_parent = current_value
                continue

            has_child_heading = any(
                rows[lower_row][column]
                for lower_row in range(
                    row_index + 1,
                    len(rows)
                )
            )

            if (
                last_parent
                and has_child_heading
            ):
                rows[row_index][column] = (
                    last_parent
                )

    combined_headers = []

    for column in range(width):
        parts = []

        for row_index in range(
            len(rows)
        ):
            part = clean_value(
                rows[row_index][column]
            )

            if not part:
                continue

            if (
                parts
                and normalize_text(part)
                == normalize_text(parts[-1])
            ):
                continue

            parts.append(part)

        combined_headers.append(
            " ".join(parts)
            or f"Column {column + 1}"
        )

    return make_unique_headers(
        combined_headers
    )


def matrix_to_records(matrix):
    # Blank rows preserve ki ja rahi hain,
    # taki repeated table blocks correctly detect hon.
    rows = [
        list(row)
        if row is not None
        else []
        for row in matrix
    ]

    if len(rows) < 2:
        return []

    width = max(
        len(row)
        for row in rows
    )

    header_index = detect_header_row(
        rows
    )

    header_end = header_index

    # Maximum three heading rows support karega.
    for index in range(
        header_index + 1,
        min(
            header_index + 3,
            len(rows)
        )
    ):
        if is_empty_row(rows[index]):
            break

        if _is_probable_header_continuation(
            rows[index]
        ):
            header_end = index
        else:
            break

    header_rows = rows[
        header_index:header_end + 1
    ]

    headers = _build_multi_row_headers(
        header_rows,
        width
    )

    # Repeated heading rows ko data banne se rokega.
    header_signatures = {
        _row_signature(row, width)
        for row in header_rows
    }

    records = []

    for source_row in rows[
        header_end + 1:
    ]:
        if is_empty_row(source_row):
            continue

        if (
            _row_signature(
                source_row,
                width
            )
            in header_signatures
        ):
            continue

        record = {}

        for index, header in enumerate(
            headers
        ):
            value = clean_value(
                source_row[index]
                if index < len(source_row)
                else ""
            )

            if value:
                record[header] = value

        if record:
            records.append(record)

    return records


def parse_xlsx(file_path):
    file_path = Path(file_path)

    if not file_path.exists():
        raise ValueError(
            f"Uploaded Excel file was not found: {file_path.name}"
        )

    if file_path.stat().st_size == 0:
        raise ValueError(
            "Uploaded Excel file is empty."
        )

    if not zipfile.is_zipfile(file_path):
        raise ValueError(
            "The uploaded file is not a valid .xlsx workbook. "
            "Open it in Microsoft Excel and use "
            "'Save As → Excel Workbook (*.xlsx)', then upload it again."
        )

    try:
        workbook = openpyxl.load_workbook(
            file_path,
            data_only=True,
            read_only=True
        )

    except zipfile.BadZipFile as error:
        raise ValueError(
            "The Excel file is damaged or incomplete. "
            "Save or download a fresh .xlsx copy."
        ) from error

    except Exception as error:
        raise ValueError(
            f"Unable to read Excel workbook: {error}"
        ) from error

    all_records = []

    try:
        for sheet in workbook.worksheets:
            matrix = []

            data_started = False
            empty_rows = 0

            maximum_rows = min(
                sheet.max_row or 1,
                5000
            )

            maximum_columns = min(
                sheet.max_column or 1,
                200
            )

            for row in sheet.iter_rows(
                min_row=1,
                max_row=maximum_rows,
                min_col=1,
                max_col=maximum_columns,
                values_only=True
            ):
                row_values = list(row)

                has_data = any(
                    clean_value(value)
                    for value in row_values
                )

                if not has_data:
                    if data_started:
                        empty_rows += 1

                        if empty_rows >= 40:
                            break

                    continue

                data_started = True
                empty_rows = 0

                while (
                    row_values
                    and not clean_value(row_values[-1])
                ):
                    row_values.pop()

                if row_values:
                    matrix.append(row_values)

            records = matrix_to_records(matrix)

            for record in records:
                record["_sheet"] = sheet.title
                all_records.append(record)

    finally:
        workbook.close()

    return all_records


def parse_xls(file_path):
    try:
        import pandas as pd

    except ImportError as error:
        raise RuntimeError(
            "Install pandas and xlrd for .xls files"
        ) from error

    workbook = pd.read_excel(
        file_path,
        sheet_name=None,
        header=None
    )

    records = []

    for sheet_name, dataframe in workbook.items():
        matrix = (
            dataframe
            .fillna("")
            .values
            .tolist()
        )

        for record in matrix_to_records(
            matrix
        ):
            record["_sheet"] = sheet_name
            records.append(record)

    return records


def parse_csv_file(file_path):
    encodings = (
        "utf-8-sig",
        "utf-8",
        "cp1252",
        "latin-1"
    )

    for encoding in encodings:
        try:
            with open(
                file_path,
                "r",
                encoding=encoding,
                newline=""
            ) as handle:
                matrix = list(
                    csv.reader(handle)
                )

                return matrix_to_records(
                    matrix
                )

        except UnicodeDecodeError:
            continue

    raise ValueError(
        "CSV encoding could not be read"
    )


def parse_key_value_blocks(lines):
    records = []
    current = {}

    for raw_line in lines:
        line = clean_value(raw_line)

        if not line:
            if current:
                records.append(current)
                current = {}

            continue

        match = re.match(
            r"^(.{1,150}?)(?:\s*[:=]\s*|\s+[–—-]\s+)(.+)$",
            line
        )

        if not match:
            continue

        key = clean_value(
            match.group(1)
        )

        value = clean_value(
            match.group(2)
        )

        if key in current and current:
            records.append(current)
            current = {}

        current[key] = value

    if current:
        records.append(current)

    return records


def parse_docx(file_path):
    if Document is None:
        raise RuntimeError(
            "Install python-docx"
        )

    document = Document(file_path)
    records = []

    for table in document.tables:
        matrix = [
            [
                clean_value(cell.text)
                for cell in row.cells
            ]
            for row in table.rows
        ]

        records.extend(
            matrix_to_records(matrix)
        )

    if records:
        return records

    lines = [
        paragraph.text
        for paragraph in document.paragraphs
    ]

    return parse_key_value_blocks(lines)


def parse_pdf(file_path):
    if fitz is None:
        raise RuntimeError(
            "Install pymupdf"
        )

    document = fitz.open(file_path)

    records = []
    lines = []

    try:
        for page in document:
            find_tables_method = getattr(
                page,
                "find_tables",
                None
            )

            if callable(find_tables_method):
                try:
                    table_finder = (
                        find_tables_method()
                    )

                    tables = getattr(
                        table_finder,
                        "tables",
                        []
                    ) or []

                    for table in tables:
                        extract_method = getattr(
                            table,
                            "extract",
                            None
                        )

                        if not callable(
                            extract_method
                        ):
                            continue

                        matrix = extract_method()

                        if matrix:
                            records.extend(
                                matrix_to_records(
                                    matrix
                                )
                            )

                except Exception as error:
                    print(
                        "PDF table detection warning:",
                        error
                    )

            text = page.get_text("text") or ""

            lines.extend(
                str(text).splitlines()
            )

    finally:
        document.close()

    if records:
        return records

    matrix = []

    for line in lines:
        columns = [
            clean_value(part)
            for part in re.split(
                r"\t+|\s{2,}",
                line
            )
            if clean_value(part)
        ]

        if len(columns) >= 2:
            matrix.append(columns)

    records = matrix_to_records(matrix)

    if records:
        return records

    return parse_key_value_blocks(lines)


def parse_structured_file(file_path):
    path = Path(file_path)
    extension = path.suffix.lower()

    if extension in {
        ".xlsx",
        ".xlsm"
    }:
        return parse_xlsx(path)

    if extension == ".xls":
        return parse_xls(path)

    if extension == ".csv":
        return parse_csv_file(path)

    if extension == ".docx":
        return parse_docx(path)

    if extension == ".pdf":
        return parse_pdf(path)

    raise ValueError(
        f"Unsupported structured file type: {extension}"
    )


def strong_field_descriptor(field):
    keys = (
        "label",
        "name",
        "id",
        "placeholder",
        "ariaLabel",
        "title",
        "autocomplete",
        "tableLabel"
    )

    return " ".join(
        str(field.get(key, "") or "")
        for key in keys
    )


def field_descriptor(field):
    keys = (
        "label",
        "name",
        "id",
        "placeholder",
        "ariaLabel",
        "title",
        "autocomplete",
        "nearText",
        "sectionTitle",
        "tableLabel"
    )

    return " ".join(
        str(field.get(key, "") or "")
        for key in keys
    )


def is_supported_field(field):
    field_type = normalize_text(
        field.get("type", "")
    )

    tag = normalize_text(
        field.get("tag", "")
    )

    role = normalize_text(
        field.get("role", "")
    )

    if field_type in {
        "hidden",
        "submit",
        "button",
        "reset",
        "password",
        "file",
        "image"
    }:
        return False

    return (
        tag in {
            "input",
            "textarea",
            "select",
            "div",
            "span"
        }
        or role in {
            "textbox",
            "combobox",
            "spinbutton"
        }
    )


def value_compatible(field, value):
    value = clean_value(value)

    if not value:
        return False

    field_type = normalize_text(
        field.get("type", "")
    )

    descriptor = normalize_text(
        field_descriptor(field)
    )

    if field_type in {
        "number",
        "range"
    }:
        return bool(
            re.search(
                r"-?\d+(?:\.\d+)?",
                value.replace(",", "")
            )
        )

    if field_type == "email":
        return bool(
            re.search(
                r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}",
                value
            )
        )

    if field_type == "tel":
        return (
            len(
                re.sub(
                    r"\D",
                    "",
                    value
                )
            )
            >= 8
        )

    if field_type == "date":
        return bool(
            re.search(
                (
                    r"\d{1,4}[./-]\d{1,2}[./-]\d{1,4}"
                    r"|"
                    r"\d{1,2}\s+[A-Za-z]+\s+\d{4}"
                ),
                value
            )
        )

    if field_type == "url":
        return bool(
            re.search(
                r"https?://|www\.|[A-Za-z0-9.-]+\.[A-Za-z]{2,}",
                value
            )
        )

    numeric_labels = (
        "amount",
        "salary",
        "percentage",
        "cgpa",
        "quantity",
        "capacity",
        "cost"
    )

    if any(
        term in descriptor
        for term in numeric_labels
    ):
        return bool(
            re.search(
                r"-?\d+(?:\.\d+)?",
                value.replace(",", "")
            )
        )

    return True


def _meaningful_tokens(text):
    normalized = normalize_text(text)

    words = re.findall(
        r"[a-z0-9]+",
        normalized
    )

    ignored_words = {
        "micro",
        "m3",
        "ug",
        "mg",
        "per",
        "cubic",
        "meter",
        "metre",
        "enter",
        "detail",
        "select",
        "option",
        "please"
    }

    return {
        word
        for word in words
        if (
            word not in ignored_words
            and not word.isdigit()
        )
    }


def _field_text_parts(field):
    important_keys = (
        "label",
        "name",
        "id",
        "placeholder",
        "ariaLabel",
        "title",
        "autocomplete",
        "tableLabel"
    )

    context_keys = (
        "nearText",
        "sectionTitle"
    )

    parts = []

    for key in important_keys:
        value = clean_value(
            field.get(key, "")
        )

        if value:
            parts.append(
                (value, 1.0)
            )

    for key in context_keys:
        value = clean_value(
            field.get(key, "")
        )

        if value:
            weight = (
                0.90
                if key == "nearText"
                else 0.75
            )

            parts.append(
                (value, weight)
            )

    return parts


def field_source_score(field, source_key):
    source_key = clean_value(
        source_key
    )

    source_normalized = normalize_text(
        source_key
    )

    source_tokens = _meaningful_tokens(
        source_key
    )

    best_score = 0.0

    for field_text, weight in _field_text_parts(
        field
    ):
        field_normalized = normalize_text(
            field_text
        )

        score = semantic_score(
            field_text,
            source_key
        )

        if (
            source_normalized in field_normalized
            or field_normalized in source_normalized
        ):
            score = max(
                score,
                0.95
            )

        field_tokens = _meaningful_tokens(
            field_text
        )

        if source_tokens and field_tokens:
            common_count = len(
                source_tokens & field_tokens
            )

            source_coverage = (
                common_count /
                len(source_tokens)
            )

            field_coverage = (
                common_count /
                len(field_tokens)
            )

            if source_tokens.issubset(
                field_tokens
            ):
                score = max(
                    score,
                    0.99
                )

            elif field_tokens.issubset(
                source_tokens
            ):
                score = max(
                    score,
                    0.96
                )

            elif common_count >= 2:
                score = max(
                    score,
                    0.55 +
                    0.40 *
                    max(
                        source_coverage,
                        field_coverage
                    )
                )

        best_score = max(
            best_score,
            score * weight
        )

    return min(
        best_score,
        1.0
    )


def map_row_to_fields(row, fields):
    candidates = []

    for field in fields:
        if not is_supported_field(field):
            continue

        for source_key, source_value in row.items():
            if str(source_key).startswith("_"):
                continue

            if not clean_value(source_value):
                continue

            score = field_source_score(
                field,
                source_key
            )

            if score < 0.60:
                continue

            if not value_compatible(
                field,
                source_value
            ):
                continue

            candidates.append((
                score,
                field,
                source_key,
                source_value
            ))

    candidates.sort(
        key=lambda item: item[0],
        reverse=True
    )

    used_fields = set()
    used_source_keys = set()
    matches = []

    for (
        score,
        field,
        source_key,
        source_value
    ) in candidates:

        field_identity = (
            field.get("fieldId")
            or field.get("index")
        )

        if field_identity in used_fields:
            continue

        if source_key in used_source_keys:
            continue

        formatted_value = format_value_for_field(
            source_value,
            field
        )

        if not formatted_value:
            continue

        confidence = int(
            round(
                min(score, 1.0) * 100
            )
        )

        matches.append({
            "index": field.get("index"),
            "fieldId": field.get("fieldId"),

            "label": clean_value(
                field.get("label")
                or field.get("placeholder")
                or field.get("name")
                or field.get("id")
                or source_key
            ),

            "raw_value": clean_value(
                source_value
            ),

            "value": clean_value(
                formatted_value
            ),

            "confidence": confidence,

            "status": (
                "auto_fill"
                if confidence >= 80
                else "review"
            ),

            "matched_key": clean_value(
                source_key
            ),

            "source":
                "dynamic_structured_mapper"
        })

        used_fields.add(
            field_identity
        )

        used_source_keys.add(
            source_key
        )

    return matches


def map_rows_to_fields(rows, fields):
    mapped_rows = []

    for index, row in enumerate(rows):
        matches = map_row_to_fields(
            row,
            fields
        )

        mapped_rows.append({
            "rowNumber": index + 1,
            "sourceRow": row,
            "matches": matches
        })

    return mapped_rows


def parse_impact_prediction_excel(file_path):
    return parse_structured_file(file_path)